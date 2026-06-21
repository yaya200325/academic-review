from __future__ import annotations

import base64
import io
import os
import re

import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

MERMAID_RENDER_URL = "https://mermaid.ink/img/{encoded}?type=png"
MERMAID_RENDER_TIMEOUT_SECONDS = 20

BODY_FONT_SIZE_PT = 12
TITLE_FONT_SIZE_PT = 16
BODY_FIRST_LINE_INDENT_PT = 24
REFERENCE_HANGING_INDENT_PT = 24
REFERENCE_CHECK_COLOR = RGBColor(0xFF, 0x00, 0x00)
CAPTION_RE = re.compile(
    r"^(?:(?:Figure|Table|Fig\.)|(?:\u56fe|\u8868))\s*\d+[\.:]?\s+\S",
    re.IGNORECASE,
)

MARKDOWN_IMAGE_RE = re.compile(r"^!\[(?P<caption>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")
MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
HTML_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)

LATEX_SYMBOLS = {
    r"\times": "×",
    r"\cdot": "·",
    r"\leq": "≤",
    r"\geq": "≥",
    r"\approx": "≈",
    r"\pm": "±",
    r"\alpha": "α",
    r"\beta": "β",
    r"\gamma": "γ",
    r"\delta": "δ",
    r"\epsilon": "ε",
    r"\lambda": "λ",
    r"\mu": "μ",
    r"\sigma": "σ",
    r"\phi": "φ",
}


def _apply_run_fonts(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(BODY_FONT_SIZE_PT)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _apply_math_run_fonts(run) -> None:
    run.font.name = "Cambria Math"
    run.font.size = Pt(BODY_FONT_SIZE_PT)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")


def _apply_caption_run_fonts(run) -> None:
    _apply_run_fonts(run)
    run.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")


def _apply_style_fonts(style, size_pt: int) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _configure_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _apply_style_fonts(normal, BODY_FONT_SIZE_PT)
    normal.paragraph_format.first_line_indent = Pt(BODY_FIRST_LINE_INDENT_PT)

    for name, size in (("Heading 1", TITLE_FONT_SIZE_PT), ("Heading 2", 14), ("Heading 3", 12)):
        style = document.styles[name]
        _apply_style_fonts(style, size)
        style.paragraph_format.first_line_indent = Pt(0)


def _add_plain_run(paragraph, text: str, *, superscript: bool = False, color: RGBColor | None = None) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    _apply_run_fonts(run)
    run.font.superscript = superscript
    if color is not None:
        run.font.color.rgb = color


def _add_caption_paragraph(document: Document, text: str) -> None:
    caption = str(text or "").strip()
    if not caption:
        return
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(caption)
    _apply_caption_run_fonts(run)


def _resolve_media_path(path: str, base_dir: str | None) -> str:
    candidate = str(path or "").strip().strip("<>").strip()
    if not candidate:
        return ""
    if os.path.isabs(candidate):
        return candidate
    if base_dir:
        return os.path.abspath(os.path.join(base_dir, candidate))
    return os.path.abspath(candidate)


def _add_local_image(document: Document, image_path: str, *, caption: str = "", base_dir: str | None = None) -> None:
    resolved_path = _resolve_media_path(image_path, base_dir)
    if not resolved_path or not os.path.exists(resolved_path):
        missing = image_path or resolved_path or "[unknown image path]"
        _add_paragraph(document, f"[image missing: {missing}]", "Normal", indent=False)
        if caption:
            _add_caption_paragraph(document, caption)
        return

    try:
        document.add_picture(resolved_path, width=Inches(5.8))
        if caption:
            _add_caption_paragraph(document, caption)
    except Exception:
        _add_paragraph(document, f"[image render failed: {image_path}]", "Normal", indent=False)
        if caption:
            _add_caption_paragraph(document, caption)


def _latex_to_unicode(text: str) -> str:
    output = str(text or "")
    for source, target in LATEX_SYMBOLS.items():
        output = output.replace(source, target)
    return output


def _extract_script_token(text: str, start: int) -> tuple[str, int]:
    if start >= len(text):
        return "", start
    if text[start] == "{":
        depth = 1
        index = start + 1
        chars: list[str] = []
        while index < len(text) and depth > 0:
            char = text[index]
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    index += 1
                    break
            if depth > 0:
                chars.append(char)
            index += 1
        return _latex_to_unicode("".join(chars)), index
    return _latex_to_unicode(text[start]), start + 1


def _append_math_runs(paragraph, expression: str) -> None:
    expr = _latex_to_unicode(expression)
    index = 0
    buffer: list[str] = []

    def flush_buffer() -> None:
        if buffer:
            run = paragraph.add_run("".join(buffer))
            _apply_math_run_fonts(run)
            buffer.clear()

    while index < len(expr):
        char = expr[index]
        if char in {"^", "_"}:
            flush_buffer()
            token, next_index = _extract_script_token(expr, index + 1)
            if token:
                run = paragraph.add_run(token)
                _apply_math_run_fonts(run)
                if char == "^":
                    run.font.superscript = True
                else:
                    run.font.subscript = True
            index = next_index
            continue
        buffer.append(char)
        index += 1
    flush_buffer()


def _append_formatted_runs(
    paragraph,
    text: str,
    *,
    superscript_citations: bool,
    parse_math: bool,
    highlight_reference_checks: bool = False,
) -> None:
    token_pattern = r"(\\\(.+?\\\)|\$.+?\$|(?:<sup>)?(?:\[\d+(?:-\d+)?\])+(?:</sup>)?)"
    if highlight_reference_checks:
        token_pattern = r"(\s*\[Reference check needed: [^\]]+\]|" + token_pattern[1:]
    pattern = re.compile(token_pattern, re.DOTALL)
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            _add_plain_run(paragraph, text[position:match.start()])
        token = match.group(0)
        if highlight_reference_checks and re.fullmatch(r"\s*\[Reference check needed: [^\]]+\]", token):
            _add_plain_run(paragraph, token, color=REFERENCE_CHECK_COLOR)
        elif superscript_citations and re.fullmatch(r"(?:<sup>)?(?:\[\d+(?:-\d+)?\])+(?:</sup>)?", token):
            _add_plain_run(
                paragraph,
                token.replace("<sup>", "").replace("</sup>", ""),
                superscript=True,
            )
        elif parse_math and ((token.startswith(r"\(") and token.endswith(r"\)")) or (token.startswith("$") and token.endswith("$"))):
            expr = token[2:-2] if token.startswith(r"\(") else token[1:-1]
            _append_math_runs(paragraph, expr)
        else:
            _add_plain_run(paragraph, token)
        position = match.end()
    if position < len(text):
        _add_plain_run(paragraph, text[position:])


def _strip_html_comments(text: str) -> str:
    return HTML_COMMENT_RE.sub("", str(text or ""))


def _strip_markdown_links(text: str) -> str:
    return MARKDOWN_LINK_RE.sub(r"\1", str(text or ""))


def process_citations(text: str) -> str:
    """
    Superscript numeric citations in body text while leaving the reference list
    numbering unchanged.
    """
    lines = _strip_html_comments(text).split("\n")
    result: list[str] = []
    in_references = False

    for line in lines:
        line = _strip_markdown_links(line)
        stripped = line.strip()
        if stripped in ["## \u53c2\u8003\u6587\u732e", "## References"]:  # Chinese heading: 参考文献
            in_references = True
            result.append(line)
            continue

        if in_references and re.match(r"^\[\d+\]", stripped):
            result.append(line)
        else:
            line = re.sub(r"\[(\d+(?:-\d+)?)\]", r"<sup>[\1]</sup>", line)
            result.append(line)

    return "\n".join(result)


def _find_residual_paper_id_citations(text: str) -> list[str]:
    candidate_pattern = re.compile(r"(?<!!)\[([^\[\]\n]+?)\](?!\()")
    residual: list[str] = []

    for raw_token in candidate_pattern.findall(text):
        token = raw_token.strip()
        if not token:
            continue
        if token.startswith("Reference check needed:"):
            continue
        if re.fullmatch(r"[A-Z]{1,2}(?:/[A-Z]{2})?", token):
            continue
        if re.fullmatch(r"\d+(?:[-,]\d+)*", token):
            continue
        residual.append(token)

    return sorted(set(residual))


def _add_paragraph(
    document: Document,
    text: str,
    style_name: str,
    *,
    indent: bool = True,
    hanging: bool = False,
    superscript_citations: bool = False,
    parse_math: bool = False,
    highlight_reference_checks: bool = False,
) -> None:
    paragraph = document.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if style_name.startswith("Heading"):
        paragraph.paragraph_format.first_line_indent = Pt(0)
    elif hanging:
        paragraph.paragraph_format.left_indent = Pt(REFERENCE_HANGING_INDENT_PT)
        paragraph.paragraph_format.first_line_indent = Pt(-REFERENCE_HANGING_INDENT_PT)
    elif indent:
        paragraph.paragraph_format.first_line_indent = Pt(BODY_FIRST_LINE_INDENT_PT)
    else:
        paragraph.paragraph_format.first_line_indent = Pt(0)
    _append_formatted_runs(
        paragraph,
        _strip_markdown_links(_strip_html_comments(text.strip())),
        superscript_citations=superscript_citations,
        parse_math=parse_math,
        highlight_reference_checks=highlight_reference_checks,
    )


def _split_inline_heading(level: int, text: str) -> tuple[str, str]:
    candidate = text.strip()
    if level == 2 and candidate.startswith("\u6458\u8981"):
        return "\u6458\u8981", candidate[len("\u6458\u8981"):].strip(" :.-")
    if level == 2 and candidate.lower().startswith("abstract"):
        return "Abstract", candidate[len("Abstract"):].strip(" :.-")
    if level in {2, 3}:
        match = re.match(r"^((?:\d+(?:\.\d+)*\.?\s+)\S.{0,80}?)(?:\s{2,}|\s+-\s+|\s+)(.+)$", candidate)
        if match:
            return match.group(1).strip(), match.group(2).strip()
    return candidate, ""


def _is_markdown_table_row(line: str) -> bool:
    text = str(line or "").strip()
    return text.startswith("|") and text.endswith("|") and text.count("|") >= 2


def _is_markdown_table_separator(line: str) -> bool:
    if not _is_markdown_table_row(line):
        return False
    cells = [cell.strip() for cell in str(line).strip().strip("|").split("|")]
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _split_markdown_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in str(line).strip().strip("|").split("|")]
    return [cell.replace("\\|", "|") for cell in cells]


def _set_table_cell_text(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.clear()
    _append_formatted_runs(
        paragraph,
        _strip_markdown_links(_strip_html_comments(str(text or "").strip())),
        superscript_citations=True,
        parse_math=False,
    )
    for run in paragraph.runs:
        run.bold = bold


def _set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, edge_data in kwargs.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def _apply_three_line_table(table) -> None:
    rows = table.rows
    if not rows:
        return
    border_on = {"val": "single", "sz": "8", "color": "000000"}
    border_off = {"val": "nil"}
    last_row_index = len(rows) - 1
    for row_index, row in enumerate(rows):
        for cell in row.cells:
            top = border_on if row_index == 0 else border_off
            bottom = border_on if row_index in {0, last_row_index} else border_off
            _set_cell_border(
                cell,
                top=top,
                bottom=bottom,
                left=border_off,
                right=border_off,
            )


def _add_markdown_table(document: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return
    header = _split_markdown_table_row(table_lines[0])
    body_lines = [line for line in table_lines[1:] if not _is_markdown_table_separator(line)]
    if not header or not body_lines:
        return

    table = document.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, cell_text in enumerate(header):
        _set_table_cell_text(table.rows[0].cells[index], cell_text, bold=True)

    for raw_line in body_lines:
        values = _split_markdown_table_row(raw_line)
        if len(values) < len(header):
            values.extend([""] * (len(header) - len(values)))
        elif len(values) > len(header):
            values = values[: len(header)]
        row_cells = table.add_row().cells
        for index, cell_text in enumerate(values):
            _set_table_cell_text(row_cells[index], cell_text)
    _apply_three_line_table(table)


def _render_mermaid_to_png(diagram_text: str) -> bytes | None:
    text = str(diagram_text or "").strip()
    if not text:
        return None
    try:
        encoded = base64.urlsafe_b64encode(text.encode("utf-8")).decode("ascii")
        response = requests.get(
            MERMAID_RENDER_URL.format(encoded=encoded),
            timeout=MERMAID_RENDER_TIMEOUT_SECONDS,
        )
        response.raise_for_status()
        if not response.content or not response.headers.get("Content-Type", "").startswith("image/"):
            return None
        return response.content
    except Exception:
        return None


def _insert_table_of_contents(document: Document, heading_text: str = "Contents") -> None:
    document.add_paragraph(heading_text, style="Heading 1").paragraph_format.first_line_indent = Pt(0)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click in Word and choose Update Field to generate the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, placeholder, end):
        run._element.append(child)
    document.add_paragraph()


def _enable_auto_field_update(document: Document) -> None:
    settings_element = document.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings_element.append(update)


def _count_headings(lines: list[str]) -> int:
    return sum(1 for line in lines if re.match(r"^#{1,6}\s+\S", line.strip()))


def _add_mermaid_diagram(document: Document, diagram_text: str) -> None:
    png_bytes = _render_mermaid_to_png(diagram_text)
    if png_bytes:
        try:
            document.add_picture(io.BytesIO(png_bytes), width=Inches(5.8))
            return
        except Exception:
            pass
    _add_paragraph(document, "[Mermaid diagram render failed; source retained below.]", "Normal", indent=False)
    for raw_line in diagram_text.splitlines():
        _add_paragraph(document, raw_line, "Normal", indent=False)


def markdown_to_docx_bytes(markdown: str, *, base_dir: str | None = None) -> bytes:
    """
    Convert markdown to docx bytes.

    Expects clean numeric citations such as [1] and [2-3]. Residual paper ID
    citations indicate that Stage 9 conversion was incomplete.
    """
    markdown = _strip_html_comments(markdown)
    residual = _find_residual_paper_id_citations(markdown)
    if residual:
        preview = ", ".join(f"[{token}]" for token in residual[:5])
        more = "" if len(residual) <= 5 else f" ... and {len(residual) - 5} more"
        raise ValueError(
            "Markdown contains unconverted paper_id citations: "
            f"{preview}{more}. Stage 9 must complete citation conversion before export."
        )

    markdown = process_citations(markdown)
    document = Document()
    _configure_document_styles(document)
    _enable_auto_field_update(document)

    in_references = False
    lines = str(markdown or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if _count_headings(lines) >= 3:
        _insert_table_of_contents(document)

    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue

        image_match = MARKDOWN_IMAGE_RE.match(line)
        if image_match:
            _add_local_image(
                document,
                image_match.group("path"),
                caption=image_match.group("caption").strip(),
                base_dir=base_dir,
            )
            index += 1
            continue

        if line.startswith("```mermaid") or line == "```mermaid":
            diagram_lines: list[str] = []
            lookahead = index + 1
            caption_line = ""
            while lookahead < len(lines):
                fence = lines[lookahead].strip()
                if fence.startswith("```"):
                    lookahead += 1
                    break
                diagram_lines.append(lines[lookahead])
                lookahead += 1
            _add_mermaid_diagram(document, "\n".join(diagram_lines).strip())
            if lookahead < len(lines):
                candidate_caption = lines[lookahead].strip()
                if CAPTION_RE.match(candidate_caption):
                    caption_line = candidate_caption
                    lookahead += 1
            if caption_line:
                _add_caption_paragraph(document, caption_line)
            index = lookahead
            continue

        if _is_markdown_table_row(line):
            if index > 0:
                candidate_caption = lines[index - 1].strip()
                if CAPTION_RE.match(candidate_caption):
                    if not document.paragraphs or document.paragraphs[-1].text.strip() != candidate_caption:
                        _add_caption_paragraph(document, candidate_caption)
            table_lines = [line]
            lookahead = index + 1
            while lookahead < len(lines):
                next_line = lines[lookahead].strip()
                if not _is_markdown_table_row(next_line):
                    break
                table_lines.append(next_line)
                lookahead += 1
            if len(table_lines) >= 2 and any(_is_markdown_table_separator(item) for item in table_lines[1:2]):
                _add_markdown_table(document, table_lines)
                index = lookahead
                continue
            for table_line in table_lines:
                _add_paragraph(document, table_line, "Normal", indent=False, parse_math=True)
            index = lookahead
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            heading_text, inline_body = _split_inline_heading(level, heading_match.group(2))
            style_name = f"Heading {level}"
            _add_paragraph(document, heading_text, style_name, indent=False)
            lower_heading = heading_text.lower()
            in_references = lower_heading in {"references", "\u53c2\u8003\u6587\u732e"}  # Chinese heading: 参考文献
            if inline_body:
                _add_paragraph(
                    document,
                    inline_body,
                    "Normal",
                    indent=not in_references,
                    hanging=in_references,
                    superscript_citations=not in_references,
                    parse_math=not in_references,
                    highlight_reference_checks=in_references,
                )
            index += 1
            continue

        if in_references:
            _add_paragraph(document, line, "Normal", indent=False, hanging=True, highlight_reference_checks=True)
        else:
            _add_paragraph(document, line, "Normal", indent=True, superscript_citations=True, parse_math=True)
        index += 1

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def main() -> None:
    import sys

    if len(sys.argv) < 3:
        print("Usage: python docx_export.py <input.md> <output.docx>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(input_path, "r", encoding="utf-8") as handle:
        markdown = handle.read()

    docx_bytes = markdown_to_docx_bytes(
        markdown,
        base_dir=os.path.dirname(os.path.abspath(input_path)),
    )

    with open(output_path, "wb") as handle:
        handle.write(docx_bytes)

    print(f"Converted {input_path} -> {output_path}")


if __name__ == "__main__":
    main()
